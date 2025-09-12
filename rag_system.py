"""
RAG (Retrieval-Augmented Generation) System for PharmBot
Supports documents, images, and text processing with ChromaDB and LangChain
"""

import streamlit as st
import chromadb
from chromadb.config import Settings
import uuid
from pathlib import Path
import json
from datetime import datetime
from typing import List, Dict, Any, Optional
import hashlib

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_community.document_loaders import (
    PyPDFLoader, TextLoader, CSVLoader, 
    UnstructuredWordDocumentLoader, UnstructuredPowerPointLoader
)
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma

# Image processing
from PIL import Image
import pytesseract
import io
import base64

# Document processing
import pandas as pd
import docx
import PyPDF2

class RAGSystem:
    """RAG System for document and image processing."""
    
    def __init__(self, user_id: str):
        self.user_id = user_id
        self.rag_data_dir = Path("user_data") / f"rag_{user_id}"
        self.rag_data_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize ChromaDB
        self.chroma_client = chromadb.PersistentClient(
            path=str(self.rag_data_dir / "chroma_db")
        )
        
        # Initialize embeddings
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}
        )
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Get or create collection
        self.collection_name = f"pharmbot_docs_{user_id}"
        try:
            self.collection = self.chroma_client.get_collection(self.collection_name)
        except:
            self.collection = self.chroma_client.create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
        
        # Document metadata file
        self.metadata_file = self.rag_data_dir / "documents_metadata.json"
        self.documents_metadata = self.load_documents_metadata()
    
    def load_documents_metadata(self) -> Dict:
        """Load documents metadata from file."""
        if self.metadata_file.exists():
            try:
                with open(self.metadata_file, 'r') as f:
                    return json.load(f)
            except:
                return {}
        return {}
    
    def save_documents_metadata(self):
        """Save documents metadata to file."""
        with open(self.metadata_file, 'w') as f:
            json.dump(self.documents_metadata, f, indent=2, default=str)
    
    def get_file_hash(self, file_content: bytes) -> str:
        """Generate hash for file content to avoid duplicates."""
        return hashlib.md5(file_content).hexdigest()
    
    def extract_text_from_image(self, image: Image.Image) -> str:
        """Extract text from image using OCR."""
        try:
            # Convert image to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using pytesseract
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            st.error(f"Error extracting text from image: {e}")
            return ""
    
    def process_pdf(self, file_content: bytes, filename: str) -> List[Document]:
        """Process PDF file and extract text."""
        documents = []
        try:
            # Save temporary file
            temp_file = self.rag_data_dir / f"temp_{filename}"
            with open(temp_file, 'wb') as f:
                f.write(file_content)
            
            # Load and process PDF
            loader = PyPDFLoader(str(temp_file))
            docs = loader.load()
            
            # Clean up temp file
            temp_file.unlink()
            
            return docs
        except Exception as e:
            st.error(f"Error processing PDF: {e}")
            return []
    
    def process_text_file(self, file_content: bytes, filename: str) -> List[Document]:
        """Process text file."""
        try:
            text = file_content.decode('utf-8')
            doc = Document(
                page_content=text,
                metadata={"source": filename, "type": "text"}
            )
            return [doc]
        except Exception as e:
            st.error(f"Error processing text file: {e}")
            return []
    
    def process_csv_file(self, file_content: bytes, filename: str) -> List[Document]:
        """Process CSV file."""
        try:
            # Read CSV
            df = pd.read_csv(io.BytesIO(file_content))
            
            # Convert to text
            text = df.to_string(index=False)
            doc = Document(
                page_content=text,
                metadata={"source": filename, "type": "csv", "rows": len(df)}
            )
            return [doc]
        except Exception as e:
            st.error(f"Error processing CSV file: {e}")
            return []
    
    def process_word_document(self, file_content: bytes, filename: str) -> List[Document]:
        """Process Word document."""
        try:
            # Save temporary file
            temp_file = self.rag_data_dir / f"temp_{filename}"
            with open(temp_file, 'wb') as f:
                f.write(file_content)
            
            # Load and process Word document
            loader = UnstructuredWordDocumentLoader(str(temp_file))
            docs = loader.load()
            
            # Clean up temp file
            temp_file.unlink()
            
            return docs
        except Exception as e:
            st.error(f"Error processing Word document: {e}")
            return []
    
    def add_document(self, file_content: bytes, filename: str, file_type: str) -> bool:
        """Add document to the RAG system."""
        try:
            # Check if document already exists
            file_hash = self.get_file_hash(file_content)
            if file_hash in self.documents_metadata:
                st.warning(f"Document '{filename}' already exists in the knowledge base.")
                return False
            
            # Process document based on type
            documents = []
            if file_type == "application/pdf":
                documents = self.process_pdf(file_content, filename)
            elif file_type == "text/plain":
                documents = self.process_text_file(file_content, filename)
            elif file_type == "text/csv":
                documents = self.process_csv_file(file_content, filename)
            elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                documents = self.process_word_document(file_content, filename)
            else:
                st.error(f"Unsupported file type: {file_type}")
                return False
            
            if not documents:
                st.error("No content extracted from the document.")
                return False
            
            # Split documents into chunks
            chunks = []
            for doc in documents:
                doc_chunks = self.text_splitter.split_documents([doc])
                chunks.extend(doc_chunks)
            
            # Add to ChromaDB
            doc_ids = []
            texts = []
            metadatas = []
            
            for i, chunk in enumerate(chunks):
                doc_id = f"{file_hash}_{i}"
                doc_ids.append(doc_id)
                texts.append(chunk.page_content)
                
                metadata = chunk.metadata.copy()
                metadata.update({
                    "filename": filename,
                    "file_hash": file_hash,
                    "chunk_index": i,
                    "added_at": datetime.now().isoformat()
                })
                metadatas.append(metadata)
            
            # Add to collection
            self.collection.add(
                ids=doc_ids,
                documents=texts,
                metadatas=metadatas
            )
            
            # Update metadata
            self.documents_metadata[file_hash] = {
                "filename": filename,
                "file_type": file_type,
                "chunks_count": len(chunks),
                "added_at": datetime.now().isoformat(),
                "size_bytes": len(file_content)
            }
            self.save_documents_metadata()
            
            return True
            
        except Exception as e:
            st.error(f"Error adding document: {e}")
            return False
    
    def add_image(self, image: Image.Image, filename: str) -> bool:
        """Add image to the RAG system by extracting text."""
        try:
            # Extract text from image
            text = self.extract_text_from_image(image)
            
            if not text:
                st.warning("No text found in the image.")
                return False
            
            # Convert image to bytes for hashing
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            img_bytes = img_byte_arr.getvalue()
            
            # Check if image already exists
            file_hash = self.get_file_hash(img_bytes)
            if file_hash in self.documents_metadata:
                st.warning(f"Image '{filename}' already exists in the knowledge base.")
                return False
            
            # Create document from extracted text
            doc = Document(
                page_content=text,
                metadata={"source": filename, "type": "image"}
            )
            
            # Split into chunks
            chunks = self.text_splitter.split_documents([doc])
            
            # Add to ChromaDB
            doc_ids = []
            texts = []
            metadatas = []
            
            for i, chunk in enumerate(chunks):
                doc_id = f"{file_hash}_{i}"
                doc_ids.append(doc_id)
                texts.append(chunk.page_content)
                
                metadata = chunk.metadata.copy()
                metadata.update({
                    "filename": filename,
                    "file_hash": file_hash,
                    "chunk_index": i,
                    "added_at": datetime.now().isoformat()
                })
                metadatas.append(metadata)
            
            # Add to collection
            self.collection.add(
                ids=doc_ids,
                documents=texts,
                metadatas=metadatas
            )
            
            # Update metadata
            self.documents_metadata[file_hash] = {
                "filename": filename,
                "file_type": "image",
                "chunks_count": len(chunks),
                "added_at": datetime.now().isoformat(),
                "size_bytes": len(img_bytes),
                "extracted_text_length": len(text)
            }
            self.save_documents_metadata()
            
            return True
            
        except Exception as e:
            st.error(f"Error adding image: {e}")
            return False
    
    def search_documents(self, query: str, n_results: int = 5) -> List[Dict]:
        """Search for relevant documents."""
        try:
            if self.collection.count() == 0:
                return []
            
            results = self.collection.query(
                query_texts=[query],
                n_results=min(n_results, self.collection.count())
            )
            
            # Format results
            formatted_results = []
            if results['documents'] and results['documents'][0]:
                for i, (doc, metadata) in enumerate(zip(results['documents'][0], results['metadatas'][0])):
                    formatted_results.append({
                        "content": doc,
                        "metadata": metadata,
                        "relevance_score": 1 - (results['distances'][0][i] if results['distances'] else 0)
                    })
            
            return formatted_results
            
        except Exception as e:
            st.error(f"Error searching documents: {e}")
            return []
    
    def get_context_for_query(self, query: str, max_context_length: int = 3000) -> str:
        """Get relevant context for a query."""
        results = self.search_documents(query, n_results=5)
        
        if not results:
            return ""
        
        context_parts = []
        current_length = 0
        
        for result in results:
            content = result["content"]
            filename = result["metadata"].get("filename", "Unknown")
            
            # Add source information
            source_info = f"\n[Source: {filename}]\n"
            full_content = source_info + content
            
            if current_length + len(full_content) <= max_context_length:
                context_parts.append(full_content)
                current_length += len(full_content)
            else:
                # Add partial content if it fits
                remaining_space = max_context_length - current_length
                if remaining_space > 100:  # Only add if meaningful space left
                    partial_content = source_info + content[:remaining_space-len(source_info)-10] + "..."
                    context_parts.append(partial_content)
                break
        
        return "\n\n".join(context_parts)
    
    def delete_document(self, file_hash: str) -> bool:
        """Delete a document from the RAG system."""
        try:
            if file_hash not in self.documents_metadata:
                return False
            
            # Get all chunk IDs for this document
            doc_info = self.documents_metadata[file_hash]
            chunk_ids = [f"{file_hash}_{i}" for i in range(doc_info["chunks_count"])]
            
            # Delete from ChromaDB
            self.collection.delete(ids=chunk_ids)
            
            # Remove from metadata
            del self.documents_metadata[file_hash]
            self.save_documents_metadata()
            
            return True
            
        except Exception as e:
            st.error(f"Error deleting document: {e}")
            return False
    
    def get_documents_list(self) -> List[Dict]:
        """Get list of all documents in the knowledge base."""
        return [
            {
                "file_hash": file_hash,
                "filename": info["filename"],
                "file_type": info["file_type"],
                "chunks_count": info["chunks_count"],
                "added_at": info["added_at"],
                "size_bytes": info.get("size_bytes", 0)
            }
            for file_hash, info in self.documents_metadata.items()
        ]
    
    def clear_all_documents(self) -> bool:
        """Clear all documents from the RAG system."""
        try:
            # Delete collection
            self.chroma_client.delete_collection(self.collection_name)
            
            # Recreate collection
            self.collection = self.chroma_client.create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            
            # Clear metadata
            self.documents_metadata = {}
            self.save_documents_metadata()
            
            return True
            
        except Exception as e:
            st.error(f"Error clearing documents: {e}")
            return False