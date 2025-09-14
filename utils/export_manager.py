"""
Export Manager for PharmGPT
Handle conversation exports to PDF and Word formats
"""

import io
import logging
from datetime import datetime
from typing import Dict, List, Optional
import streamlit as st

# Configure logging
logger = logging.getLogger(__name__)

class ConversationExporter:
    """Handle conversation exports to various formats."""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt', 'md']
    
    def export_conversation(self, conversation_data: Dict, format_type: str = 'pdf') -> Optional[bytes]:
        """Export conversation to specified format."""
        try:
            if format_type not in self.supported_formats:
                raise ValueError(f"Unsupported format: {format_type}")
            
            if format_type == 'pdf':
                return self._export_to_pdf(conversation_data)
            elif format_type == 'docx':
                return self._export_to_docx(conversation_data)
            elif format_type == 'txt':
                return self._export_to_txt(conversation_data)
            elif format_type == 'md':
                return self._export_to_markdown(conversation_data)
            
        except Exception as e:
            logger.error(f"Export failed for format {format_type}: {e}")
            return None
    
    def _export_to_pdf(self, conversation_data: Dict) -> bytes:
        """Export conversation to PDF format."""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.colors import HexColor
            
            # Create PDF buffer
            buffer = io.BytesIO()
            
            # Create document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=HexColor('#2563eb'),
                alignment=1  # Center alignment
            )
            
            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=20,
                textColor=HexColor('#64748b'),
                alignment=1
            )
            
            user_style = ParagraphStyle(
                'UserMessage',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                leftIndent=20,
                rightIndent=20,
                borderColor=HexColor('#e0f2fe'),
                borderWidth=1,
                borderPadding=10,
                backColor=HexColor('#f0f9ff')
            )
            
            assistant_style = ParagraphStyle(
                'AssistantMessage',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                leftIndent=20,
                rightIndent=20,
                borderColor=HexColor('#dcfce7'),
                borderWidth=1,
                borderPadding=10,
                backColor=HexColor('#f0fdf4')
            )
            
            timestamp_style = ParagraphStyle(
                'Timestamp',
                parent=styles['Normal'],
                fontSize=10,
                textColor=HexColor('#6b7280'),
                spaceAfter=6
            )
            
            # Build document content
            story = []
            
            # Title and metadata
            story.append(Paragraph("💊 PharmGPT Conversation Export", title_style))
            story.append(Spacer(1, 12))
            
            # Conversation info
            title = conversation_data.get('title', 'Untitled Conversation')
            export_date = datetime.now().strftime('%B %d, %Y at %I:%M %p')
            
            story.append(Paragraph(f"<b>Conversation:</b> {title}", subtitle_style))
            story.append(Paragraph(f"<b>Exported:</b> {export_date}", subtitle_style))
            story.append(Paragraph(f"<b>Messages:</b> {len(conversation_data.get('messages', []))}", subtitle_style))
            
            # Handle bulk export info
            if conversation_data.get('export_type') == 'bulk':
                story.append(Paragraph(f"<b>Total Conversations:</b> {conversation_data.get('conversation_count', 0)}", subtitle_style))
                story.append(Paragraph(f"<b>Export Type:</b> Bulk Export", subtitle_style))
            else:
                # Model info for single conversations
                model = conversation_data.get('model', 'normal')
                model_name = "Turbo Mode (Sonoma Sky Alpha)" if model == 'turbo' else "Normal Mode (Llama 4 Maverick)"
                story.append(Paragraph(f"<b>AI Model:</b> {model_name}", subtitle_style))
            
            story.append(Spacer(1, 30))
            
            # Messages
            messages = conversation_data.get('messages', [])
            
            for i, message in enumerate(messages):
                role = message.get('role', 'unknown')
                content = message.get('content', '')
                timestamp = message.get('timestamp', '')
                
                # Format timestamp
                if timestamp:
                    try:
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        formatted_time = dt.strftime('%I:%M %p')
                    except:
                        formatted_time = timestamp
                else:
                    formatted_time = f"Message {i+1}"
                
                # Add timestamp
                story.append(Paragraph(f"🕒 {formatted_time}", timestamp_style))
                
                # Add message based on role
                if role == 'user':
                    # Clean and escape content for PDF
                    clean_content = self._clean_content_for_pdf(content)
                    story.append(Paragraph(f"<b>👤 You:</b><br/>{clean_content}", user_style))
                elif role == 'assistant':
                    clean_content = self._clean_content_for_pdf(content)
                    story.append(Paragraph(f"<b>🤖 PharmGPT:</b><br/>{clean_content}", assistant_style))
                elif role == 'system':
                    # Handle conversation separators for bulk exports
                    if 'CONVERSATION:' in content or 'END OF CONVERSATION' in content:
                        separator_style = ParagraphStyle(
                            'Separator',
                            parent=styles['Heading3'],
                            fontSize=14,
                            spaceAfter=12,
                            spaceBefore=12,
                            textColor=HexColor('#7c3aed'),
                            alignment=1,  # Center alignment
                            borderColor=HexColor('#7c3aed'),
                            borderWidth=1,
                            borderPadding=8
                        )
                        story.append(Paragraph(content, separator_style))
                
                story.append(Spacer(1, 12))
            
            # Footer
            story.append(Spacer(1, 30))
            story.append(Paragraph(
                "<i>This conversation was exported from PharmGPT - AI Pharmacology Assistant<br/>"
                "For educational purposes only. Always consult healthcare professionals for clinical decisions.</i>",
                styles['Normal']
            ))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            return pdf_bytes
            
        except Exception as e:
            logger.error(f"PDF export error: {e}")
            raise
    
    def _export_to_docx(self, conversation_data: Dict) -> bytes:
        """Export conversation to Word document format."""
        try:
            from docx import Document
            from docx.shared import Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            
            # Create document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Title
            title = doc.add_heading('💊 PharmGPT Conversation Export', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Conversation metadata
            doc.add_paragraph()
            
            # Conversation info table
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Shading Accent 1'
            
            # Fill table with metadata
            conv_title = conversation_data.get('title', 'Untitled Conversation')
            export_date = datetime.now().strftime('%B %d, %Y at %I:%M %p')
            message_count = len(conversation_data.get('messages', []))
            model = conversation_data.get('model', 'normal')
            model_name = "Turbo Mode (Sonoma Sky Alpha)" if model == 'turbo' else "Normal Mode (Llama 4 Maverick)"
            
            table.cell(0, 0).text = 'Conversation Title'
            table.cell(0, 1).text = conv_title
            table.cell(1, 0).text = 'Export Date'
            table.cell(1, 1).text = export_date
            table.cell(2, 0).text = 'Total Messages'
            table.cell(2, 1).text = str(message_count)
            table.cell(3, 0).text = 'AI Model Used'
            table.cell(3, 1).text = model_name
            
            doc.add_paragraph()
            doc.add_paragraph()
            
            # Messages
            messages = conversation_data.get('messages', [])
            
            for i, message in enumerate(messages):
                role = message.get('role', 'unknown')
                content = message.get('content', '')
                timestamp = message.get('timestamp', '')
                
                # Format timestamp
                if timestamp:
                    try:
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        formatted_time = dt.strftime('%I:%M %p')
                    except:
                        formatted_time = timestamp
                else:
                    formatted_time = f"Message {i+1}"
                
                # Add timestamp
                time_para = doc.add_paragraph(f"🕒 {formatted_time}")
                time_para.style = 'Intense Quote'
                
                # Add message based on role
                if role == 'user':
                    # User message
                    user_para = doc.add_paragraph()
                    user_run = user_para.add_run(f"👤 You: ")
                    user_run.bold = True
                    user_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue color
                    
                    content_run = user_para.add_run(content)
                    
                    # Add shading to user messages
                    self._add_paragraph_shading(user_para, RGBColor(240, 249, 255))
                    
                elif role == 'assistant':
                    # Assistant message
                    assistant_para = doc.add_paragraph()
                    assistant_run = assistant_para.add_run(f"🤖 PharmGPT: ")
                    assistant_run.bold = True
                    assistant_run.font.color.rgb = RGBColor(34, 197, 94)  # Green color
                    
                    content_run = assistant_para.add_run(content)
                    
                    # Add shading to assistant messages
                    self._add_paragraph_shading(assistant_para, RGBColor(240, 253, 244))
                
                doc.add_paragraph()
            
            # Footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph(
                "This conversation was exported from PharmGPT - AI Pharmacology Assistant\n"
                "For educational purposes only. Always consult healthcare professionals for clinical decisions."
            )
            footer_para.style = 'Intense Quote'
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            
            # Get document bytes
            docx_bytes = buffer.getvalue()
            buffer.close()
            
            return docx_bytes
            
        except Exception as e:
            logger.error(f"DOCX export error: {e}")
            raise
    
    def _export_to_txt(self, conversation_data: Dict) -> bytes:
        """Export conversation to plain text format."""
        try:
            lines = []
            
            # Header
            lines.append("=" * 60)
            lines.append("PharmGPT Conversation Export")
            lines.append("=" * 60)
            lines.append("")
            
            # Metadata
            title = conversation_data.get('title', 'Untitled Conversation')
            export_date = datetime.now().strftime('%B %d, %Y at %I:%M %p')
            message_count = len(conversation_data.get('messages', []))
            model = conversation_data.get('model', 'normal')
            model_name = "Turbo Mode" if model == 'turbo' else "Normal Mode"
            
            lines.append(f"Conversation: {title}")
            lines.append(f"Exported: {export_date}")
            lines.append(f"Messages: {message_count}")
            lines.append(f"AI Model: {model_name}")
            lines.append("")
            lines.append("-" * 60)
            lines.append("")
            
            # Messages
            messages = conversation_data.get('messages', [])
            
            for i, message in enumerate(messages):
                role = message.get('role', 'unknown')
                content = message.get('content', '')
                timestamp = message.get('timestamp', '')
                
                # Format timestamp
                if timestamp:
                    try:
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        formatted_time = dt.strftime('%I:%M %p')
                    except:
                        formatted_time = timestamp
                else:
                    formatted_time = f"Message {i+1}"
                
                lines.append(f"[{formatted_time}]")
                
                if role == 'user':
                    lines.append(f"YOU: {content}")
                elif role == 'assistant':
                    lines.append(f"PHARMGPT: {content}")
                
                lines.append("")
            
            # Footer
            lines.append("-" * 60)
            lines.append("This conversation was exported from PharmGPT")
            lines.append("For educational purposes only.")
            lines.append("Always consult healthcare professionals for clinical decisions.")
            
            # Join and encode
            text_content = "\n".join(lines)
            return text_content.encode('utf-8')
            
        except Exception as e:
            logger.error(f"TXT export error: {e}")
            raise
    
    def _export_to_markdown(self, conversation_data: Dict) -> bytes:
        """Export conversation to Markdown format."""
        try:
            lines = []
            
            # Header
            title = conversation_data.get('title', 'Untitled Conversation')
            lines.append(f"# 💊 {title}")
            lines.append("")
            lines.append("*Exported from PharmGPT - AI Pharmacology Assistant*")
            lines.append("")
            
            # Metadata
            export_date = datetime.now().strftime('%B %d, %Y at %I:%M %p')
            message_count = len(conversation_data.get('messages', []))
            model = conversation_data.get('model', 'normal')
            model_name = "Turbo Mode (Sonoma Sky Alpha)" if model == 'turbo' else "Normal Mode (Llama 4 Maverick)"
            
            lines.append("## Conversation Details")
            lines.append("")
            lines.append(f"- **Exported:** {export_date}")
            lines.append(f"- **Messages:** {message_count}")
            lines.append(f"- **AI Model:** {model_name}")
            lines.append("")
            lines.append("---")
            lines.append("")
            
            # Messages
            messages = conversation_data.get('messages', [])
            
            for i, message in enumerate(messages):
                role = message.get('role', 'unknown')
                content = message.get('content', '')
                timestamp = message.get('timestamp', '')
                
                # Format timestamp
                if timestamp:
                    try:
                        dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                        formatted_time = dt.strftime('%I:%M %p')
                    except:
                        formatted_time = timestamp
                else:
                    formatted_time = f"Message {i+1}"
                
                if role == 'user':
                    lines.append(f"### 👤 You _{formatted_time}_")
                    lines.append("")
                    lines.append(f"> {content}")
                    lines.append("")
                elif role == 'assistant':
                    lines.append(f"### 🤖 PharmGPT _{formatted_time}_")
                    lines.append("")
                    lines.append(content)
                    lines.append("")
            
            # Footer
            lines.append("---")
            lines.append("")
            lines.append("*This conversation was exported from PharmGPT - AI Pharmacology Assistant*")
            lines.append("")
            lines.append("**Disclaimer:** For educational purposes only. Always consult healthcare professionals for clinical decisions.")
            
            # Join and encode
            markdown_content = "\n".join(lines)
            return markdown_content.encode('utf-8')
            
        except Exception as e:
            logger.error(f"Markdown export error: {e}")
            raise
    
    def _clean_content_for_pdf(self, content: str) -> str:
        """Clean content for PDF export."""
        # Escape special characters and handle formatting
        content = content.replace('&', '&amp;')
        content = content.replace('<', '&lt;')
        content = content.replace('>', '&gt;')
        
        # Convert markdown-style formatting to HTML
        content = content.replace('**', '<b>')
        content = content.replace('*', '<i>')
        
        # Handle line breaks
        content = content.replace('\n', '<br/>')
        
        return content
    
    def _add_paragraph_shading(self, paragraph, color):
        """Add background shading to a paragraph in Word document."""
        try:
            from docx.oxml.shared import OxmlElement, qn
            
            # Create shading element
            shading_elm = OxmlElement(qn('w:shd'))
            shading_elm.set(qn('w:val'), 'clear')
            shading_elm.set(qn('w:color'), 'auto')
            shading_elm.set(qn('w:fill'), f"{color.r:02x}{color.g:02x}{color.b:02x}")
            
            # Add to paragraph properties
            paragraph._element.get_or_add_pPr().append(shading_elm)
        except:
            # If shading fails, continue without it
            pass
    
    def get_filename(self, conversation_data: Dict, format_type: str) -> str:
        """Generate appropriate filename for export."""
        title = conversation_data.get('title', 'PharmGPT_Conversation')
        
        # Clean title for filename
        clean_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        clean_title = clean_title.replace(' ', '_')
        
        # Limit length
        if len(clean_title) > 50:
            clean_title = clean_title[:50]
        
        # Add timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        return f"PharmGPT_{clean_title}_{timestamp}.{format_type}"

# Global exporter instance
conversation_exporter = ConversationExporter()

def export_conversation_to_format(conversation_data: Dict, format_type: str) -> Optional[bytes]:
    """Export conversation to specified format."""
    return conversation_exporter.export_conversation(conversation_data, format_type)

def get_export_filename(conversation_data: Dict, format_type: str) -> str:
    """Get appropriate filename for export."""
    return conversation_exporter.get_filename(conversation_data, format_type)